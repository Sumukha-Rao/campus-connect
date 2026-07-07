# -*- coding: utf-8 -*-
"""Build the final College Forum App report: front pages (per format),
heading styles + auto TOC, real screenshots, accuracy edits, references."""
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC   = r"D:/MCA/MAD/Project/campus-connect/Mad_Report_EL (1).docx"
OUT   = r"D:/MCA/MAD/Project/campus-connect/College_Forum_App_Report.docx"
SHOTS = r"D:/MCA/MAD/Project/campus-connect/report_assets/shots"
LOGO  = r"D:/MCA/MAD/Project/campus-connect/report_assets/format_media/image1.png"
FONT  = "Bookman Old Style"

TITLE = "College Forum App"
STUDENTS = [  # alphabetical
    ("Priya N", "1RV25MC075"),
    ("Spandana N", "1RV25MC096"),
    ("Sumanth Bhaskar Hegde", "1RV25MC103"),
    ("Sumukha Rao B S", "1RV25MC104"),
]

d = Document(SRC)
body = d.element.body
ref = d.paragraphs[0]._p           # "1. Introduction" – front matter goes before this
parent = d.paragraphs[0]._parent

# ---------- helpers ----------
def set_font(run, size=12, bold=False, italic=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    for a in ('w:ascii','w:hAnsi','w:cs'):
        rFonts.set(qn(a), font)

def new_par_before(refel, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=0):
    p = OxmlElement('w:p')
    refel.addprevious(p)
    para = Paragraph(p, parent)
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    return para

def text_par_before(refel, text="", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=2, space_before=0, italic=False, color=None):
    para = new_par_before(refel, align, space_after, space_before)
    if text:
        r = para.add_run(text); set_font(r, size, bold, italic, color)
    return para

def img_par_before(refel, path, width_in, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4):
    para = new_par_before(refel, align, space_after)
    r = para.add_run(); r.add_picture(path, width=Inches(width_in))
    return para

def page_break_before(refel):
    para = new_par_before(refel, WD_ALIGN_PARAGRAPH.LEFT, 0)
    r = para.add_run(); br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); r._r.append(br)
    return para

def set_cell(cell, lines, size=12, bold_first=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    if isinstance(lines, str): lines = [lines]
    for i, ln in enumerate(lines):
        para = p if i == 0 else cell.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(ln); set_font(r, size, bold=(bold_first and i == 0))

def borderless(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge); e.set(qn('w:val'),'none'); borders.append(e)
    tblPr.append(borders)

def bordered(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000'); borders.append(e)
    tblPr.append(borders)

def move_table_before(tbl, refel):
    refel.addprevious(tbl._tbl)

def make_table(rows, cols, widths=None):
    t = d.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for r in t.rows:
            for c, w in zip(r.cells, widths):
                c.width = Inches(w)
    return t

# ============================================================
# COVER PAGE
# ============================================================
img_par_before(ref, LOGO, 4.6, space_after=6)
text_par_before(ref, "MOBILE APPLICATION DEVELOPMENT", 14, True, space_after=0)
text_par_before(ref, "MCA221IA", 14, True, space_after=0)
text_par_before(ref, "SEE PROJECT BASED LABORATORY", 14, True, space_after=10)
text_par_before(ref, TITLE, 16, True, space_after=10, space_before=6)
text_par_before(ref, "submitted by", 12, False, space_after=6)

# student table (Name | USN)
stbl = make_table(len(STUDENTS), 2, widths=[3.6, 2.0])
borderless(stbl)
for i,(nm,usn) in enumerate(STUDENTS):
    set_cell(stbl.rows[i].cells[0], nm, 12, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(stbl.rows[i].cells[1], usn, 12, align=WD_ALIGN_PARAGRAPH.LEFT)
move_table_before(stbl, ref)

text_par_before(ref, "under the guidance of", 12, False, space_after=4, space_before=8)
gtbl = make_table(1, 1, widths=[5.0]); borderless(gtbl)
set_cell(gtbl.rows[0].cells[0],
         ["Prof. Prashanth K", "Assistant Professor", "Department of MCA", "RV College of Engineering"],
         12, align=WD_ALIGN_PARAGRAPH.CENTER)
move_table_before(gtbl, ref)

text_par_before(ref, "Department of Master of Computer Applications", 12, True, space_after=0, space_before=10)
text_par_before(ref, "RV College of Engineering, Bengaluru – 560059", 12, True, space_after=0)
text_par_before(ref, "2025-2026", 12, True, space_after=0)
page_break_before(ref)

# ============================================================
# CERTIFICATE PAGE
# ============================================================
img_par_before(ref, LOGO, 4.6, space_after=8)
text_par_before(ref, "CERTIFICATE", 13, True, space_after=10)

names_join = ", ".join(f"{nm} ({usn})" for nm,usn in STUDENTS[:-1]) + \
             f" and {STUDENTS[-1][0]} ({STUDENTS[-1][1]})"
cpara = new_par_before(ref, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
cpara.paragraph_format.line_spacing = 1.5
def addr(txt, bold=False):
    r = cpara.add_run(txt); set_font(r, 12, bold)
addr("Certified that the project entitled ")
addr(f"“{TITLE}”", True)
addr(" on ")
addr("Mobile Application Development – MCA221IA", True)
addr(" has been carried out by ")
addr(names_join, True)
addr(" who have successfully completed the project for the final SEE Lab Examination, "
     "incorporating all concepts of the course conducted by the Department of MCA, "
     "RV College of Engineering, Bengaluru.")

# guide / HOD table
ctbl = make_table(1, 2, widths=[3.0, 3.0]); borderless(ctbl)
set_cell(ctbl.rows[0].cells[0],
         ["Internal Guide", "Prof. Prashanth K", "Assistant Professor", "Department of MCA",
          "RV College of Engineering"], 12, align=WD_ALIGN_PARAGRAPH.LEFT)
set_cell(ctbl.rows[0].cells[1],
         ["Head of the Department", "Dr. Jasmine K S", "Associate Professor & Director",
          "Department of MCA,", "RV College of Engineering"], 12, align=WD_ALIGN_PARAGRAPH.RIGHT)
move_table_before(ctbl, ref)

text_par_before(ref, "External Viva Examination", 12, True, space_after=6, space_before=16)
ep = new_par_before(ref, WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
ep.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.LEFT)
r = ep.add_run("Name of Examiners\t\t\tSignature with Date"); set_font(r, 12, True)
for label in ("1.", "2."):
    pp = new_par_before(ref, WD_ALIGN_PARAGRAPH.LEFT, space_after=16)
    r = pp.add_run(label); set_font(r, 12)
page_break_before(ref)

# ============================================================
# TABLE OF CONTENTS PAGE  (auto field)
# ============================================================
text_par_before(ref, "Table of Contents", 16, True, space_after=10)
tocp = new_par_before(ref, WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
run = tocp.add_run()
fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'),'begin')
it = OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text = 'TOC \\o "1-2" \\h \\z \\u'
fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'),'separate')
tt = OxmlElement('w:t'); tt.text = 'Right-click here and choose “Update Field” to build the Table of Contents.'
fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'),'end')
for el in (fb, it, fs, tt, fe): run._r.append(el)
set_font(run, 12)

# section break (front matter = its own section, no header/footer)
final_sectPr = body.find(qn('w:sectPr'))
sb = new_par_before(ref, WD_ALIGN_PARAGRAPH.LEFT, 0)
pPr = sb._p.get_or_add_pPr()
sect = copy.deepcopy(final_sectPr)
for tag in ('w:headerReference','w:footerReference'):
    for el in sect.findall(qn(tag)): sect.remove(el)
pPr.append(sect)

import re

# ============================================================
# HEADING STYLES + restyle (enables auto TOC)
# ============================================================
for name, size in (('Heading 1', 15), ('Heading 2', 13)):
    st = d.styles[name]
    st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0); st.font.name = 'Times New Roman'
    st.paragraph_format.space_before = Pt(12 if name == 'Heading 1' else 8)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True

for p in list(d.paragraphs):
    t = p.text.strip()
    if re.match(r'^\d+\.\s+\S', t):
        p.style = d.styles['Heading 1']
    elif re.match(r'^\d+\.\d', t):
        p.style = d.styles['Heading 2']

# ============================================================
# ACCURACY EDITS – align report with the real app
# ============================================================
def append_run(par, text, size=11, italic=False, bold=False, font='Times New Roman'):
    r = par.add_run(text); set_font(r, size, bold, italic, font=font); return r

for p in d.paragraphs:
    t = p.text.strip()
    if t.startswith("College Forum App is a Progressive Web Application"):
        append_run(p, " Within its interface the application is branded as “RVCE Connect”, the "
                      "official notice board of RV College of Engineering.")
    elif t.startswith("To improve user engagement"):
        append_run(p, " It also delivers real-time browser push notifications and lets every "
                      "department and club operate as a subscribable community.")

# new subsections 2.9 / 2.10 before "Conceptual Feature Flow"
def find_par(text_starts):
    for p in d.paragraphs:
        if p.text.strip().startswith(text_starts):
            return p
    return None

def insert_par_before(el, text, style=None, size=11, bold=False,
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY, font='Times New Roman'):
    np = OxmlElement('w:p'); el.addprevious(np)
    para = Paragraph(np, parent); para.alignment = align
    if style: para.style = d.styles[style]
    if text:
        r = para.add_run(text)
        if style:                       # match the original headings' explicit TNR runs
            hs = 15 if style == 'Heading 1' else 13
            set_font(r, hs, True, color=(0, 0, 0), font='Times New Roman')
        else:
            set_font(r, size, bold, font=font)
    return para

cf = find_par("Conceptual Feature Flow")
if cf is not None:
    cf_el = cf._p
    insert_par_before(cf_el, "2.9 Real-Time Push Notifications", style='Heading 2',
                      align=WD_ALIGN_PARAGRAPH.LEFT)
    p29 = insert_par_before(cf_el, "Beyond the in-app feed, College Forum App supports genuine "
        "browser push notifications using the Web Push protocol with VAPID keys. For every "
        "community a viewer can switch on a per-community bell; when a publisher of that community "
        "posts a new notice, the server fans out a push message so that subscribers are alerted "
        "even when the application is closed. This keeps time-critical announcements such as "
        "placement drives and last-minute venue changes from being missed.")
    append_run(p29, " [9]")
    insert_par_before(cf_el, "2.10 Post Scheduling and Automatic Expiry", style='Heading 2',
                      align=WD_ALIGN_PARAGRAPH.LEFT)
    insert_par_before(cf_el, "While composing an announcement a publisher may set an optional "
        "“Expires On” date and an optional scheduled publish time. A background job runs every "
        "fifteen minutes and moves expired notices into a separate archive table so that feeds "
        "stay relevant and uncluttered, while administrators retain a read-only view of archived "
        "posts for historical reference. Deletion of posts is restricted to administrators and "
        "every deletion is recorded in an audit log.")

# ============================================================
# IN-TEXT CITATIONS
# ============================================================
CITE = [
    ("Node.js serves as the runtime", "[1]"),
    ("Express.js is used as the backend", "[2]"),
    ("MySQL is used as the relational database management system", "[3]"),
    ("JWT is used to implement secure authentication", "[6]"),
    ("Service Workers enable offline functionality", "[5]"),
    ("The Web App Manifest defines", "[8]"),
    ("Bootstrap 5 is utilized as the primary frontend framework", "[7]"),
    ("npm is used to manage project dependencies", "[10]"),
    ("College Forum is implemented as a Progressive Web Application", "[4]"),
]
for p in d.paragraphs:
    t = p.text.strip()
    for sw, c in CITE:
        if t.startswith(sw):
            append_run(p, " " + c)
            break

# ============================================================
# SECTION 5 – replace wireframe placeholders with real screenshots
# ============================================================
SEC5 = {
    '5.1 Login Screen': ('01_login.png', 5.3, 'Figure 5.1: Login screen of the College Forum App'),
    '5.2 Student (Viewer) Dashboard': ('03_viewer_feed.png', 5.0, 'Figure 5.2: Student (viewer) feed dashboard'),
    '5.3 Publisher Dashboard': ('05_publisher_compose.png', 5.0, 'Figure 5.3: Publisher create-announcement screen'),
    '5.4 Admin Dashboard': ('07_admin_dashboard.png', 5.4, 'Figure 5.4: Admin dashboard – member directory and community management'),
    '5.5 Department Management Screen': ('04_viewer_communities.png', 5.2, 'Figure 5.5: Communities directory (departments and clubs)'),
    '5.6 Publisher Subscription Screen': ('11_mobile_communities.png', 2.5, 'Figure 5.6: Subscribing to a community'),
    '5.7 Offline Page': ('09_offline.png', 5.2, 'Figure 5.7: Offline fallback page'),
    '5.8 Mobile Navigation Layout': ('10_mobile_feed.png', 2.5, 'Figure 5.8: Mobile-first responsive layout'),
}

def add_caption_after(el, text):
    cp = OxmlElement('w:p'); el.addnext(cp)
    para = Paragraph(cp, parent); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2); para.paragraph_format.space_after = Pt(8)
    r = para.add_run(text); set_font(r, 10.5, italic=True, font='Times New Roman')
    return cp

def set_image_paragraph(el, path, width):
    for rr in el.findall(qn('w:r')): el.remove(rr)
    para = Paragraph(el, parent); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    rn = para.add_run(); rn.add_picture(path, width=Inches(width))

cur = None
jobs = []
for p in d.paragraphs:
    t = p.text.strip()
    if t in SEC5:
        cur = t
    elif t == 'Wireframe' and cur:
        jobs.append((p, cur)); cur = None

for wf, key in jobs:
    fname, w, cap = SEC5[key]
    path = SHOTS + '/' + fname
    wf_el = wf._p
    img_el = None
    sib = wf_el.getnext(); steps = 0
    while sib is not None and steps < 6:
        if sib.tag == qn('w:p'):
            pr = Paragraph(sib, parent)
            if pr.style.name.startswith('Heading'):
                break
            if sib.find('.//' + qn('w:drawing')) is not None:
                img_el = sib; break
            if pr.text.strip() in ('Purpose',) or pr.text.strip() in SEC5:
                break
        sib = sib.getnext(); steps += 1
    if img_el is not None:
        set_image_paragraph(img_el, path, w)
        anchor = img_el
    else:
        np = OxmlElement('w:p'); wf_el.addnext(np)
        set_image_paragraph(np, path, w)
        anchor = np
    add_caption_after(anchor, cap)
    for rr in wf_el.findall(qn('w:r')): wf_el.remove(rr)   # clear the 'Wireframe' label

# ============================================================
# SECTION 7 – demonstration screenshots
# ============================================================
demo = find_par("7.2 Demonstration Flow")
if demo is not None:
    demo_el = demo._p
    for fname, wdt, cap in (
        ('02_register.png', 5.0, 'Figure 7.1: New-student self-registration'),
        ('06_publisher_feed.png', 5.0, 'Figure 7.2: Campus feed after a publisher posts an announcement'),
    ):
        ip = OxmlElement('w:p'); demo_el.addprevious(ip)
        set_image_paragraph(ip, SHOTS + '/' + fname, wdt)
        cp = OxmlElement('w:p'); ip.addnext(cp)
        para = Paragraph(cp, parent); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(8)
        r = para.add_run(cap); set_font(r, 10.5, italic=True, font='Times New Roman')

# caption for the existing demonstration-flow diagram
for p in d.paragraphs:
    if p._p.find('.//' + qn('w:drawing')) is not None and p.text.strip() == '':
        # only the demo-flow diagram lives just after the 7.2 heading region; skip others
        pass

# ============================================================
# REFERENCES
# ============================================================
REFS = [
    "[1] OpenJS Foundation, “Node.js Documentation.” [Online]. Available: https://nodejs.org/en/docs",
    "[2] OpenJS Foundation, “Express – Fast, unopinionated, minimalist web framework for Node.js.” [Online]. Available: https://expressjs.com",
    "[3] Oracle Corporation, “MySQL 8.0 Reference Manual.” [Online]. Available: https://dev.mysql.com/doc/refman/8.0/en/",
    "[4] Mozilla, “Progressive web apps (PWAs),” MDN Web Docs. [Online]. Available: https://developer.mozilla.org/en-US/docs/Web/Progressive_web_apps",
    "[5] Mozilla, “Service Worker API,” MDN Web Docs. [Online]. Available: https://developer.mozilla.org/en-US/docs/Web/API/Service_Worker_API",
    "[6] M. Jones, J. Bradley and N. Sakimura, “JSON Web Token (JWT),” RFC 7519, IETF, May 2015. [Online]. Available: https://datatracker.ietf.org/doc/html/rfc7519",
    "[7] Bootstrap Team, “Bootstrap 5 Documentation.” [Online]. Available: https://getbootstrap.com/docs/5.3/",
    "[8] W3C, “Web Application Manifest.” [Online]. Available: https://www.w3.org/TR/appmanifest/",
    "[9] M. Thomson, “Voluntary Application Server Identification (VAPID) for Web Push,” RFC 8292, IETF, Nov. 2017. [Online]. Available: https://datatracker.ietf.org/doc/html/rfc8292",
    "[10] npm, Inc., “npm Documentation.” [Online]. Available: https://docs.npmjs.com",
    "[11] J. Wilander et al., “bcrypt.js – Optimized bcrypt in JavaScript.” [Online]. Available: https://github.com/dcodeIO/bcrypt.js",
]
refh = d.add_paragraph()
refh.style = d.styles['Heading 1']
refh.paragraph_format.page_break_before = True
set_font(refh.add_run("9. References"), 15, True, color=(0, 0, 0), font='Times New Roman')
for entry in REFS:
    rp = d.add_paragraph()
    rp.paragraph_format.space_after = Pt(4)
    rp.paragraph_format.left_indent = Inches(0.3)
    rp.paragraph_format.first_line_indent = Inches(-0.3)
    r = rp.add_run(entry); set_font(r, 10.5, font='Times New Roman')

# ============================================================
# HEADER / FOOTER + page numbering on the body section
# ============================================================
sec = d.sections[-1]
sec_pr = sec._sectPr
pg = OxmlElement('w:pgNumType'); pg.set(qn('w:start'), '1'); sec_pr.append(pg)

def page_field(par, prefix=""):
    if prefix:
        r = par.add_run(prefix); set_font(r, 10, font='Times New Roman')
    run = par.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    for el in (fb, it, fe): run._r.append(el)
    set_font(run, 10, font='Times New Roman')

hdr = sec.header; hdr.is_linked_to_previous = False
hp = hdr.paragraphs[0]; hp.text = ''
hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
r = hp.add_run(TITLE); set_font(r, 10, font='Times New Roman')
r = hp.add_run("\tJuly 2026"); set_font(r, 10, font='Times New Roman')

ftr = sec.footer; ftr.is_linked_to_previous = False
fp = ftr.paragraphs[0]; fp.text = ''
fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
r = fp.add_run("Department of MCA, RVCE"); set_font(r, 10, font='Times New Roman')
r = fp.add_run("\t"); set_font(r, 10, font='Times New Roman')
page_field(fp, prefix="Page ")

# ============================================================
# Make Word refresh fields (TOC + page numbers) on open
# ============================================================
settings = d.settings.element
uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
settings.insert(0, uf)

d.save(OUT)
print("REPORT BUILT ->", OUT)
print("section-5 figures inserted:", len(jobs))

